"""
快速转换 Markdown → .doc → .pdf
"""
import os, sys, re, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = r"D:\AIFireCheck\combined_system\报告模板"
MD = os.path.join(BASE, "课程设计报告.md")
DOC = os.path.join(BASE, "课程设计报告.doc")
PDF = os.path.join(BASE, "课程设计报告.pdf")
IMG = os.path.join(BASE, "截图或图片")

def sf(run, n='宋体', s=12, b=False, c=None):
    run.font.name = n; run._element.rPr.rFonts.set(qn('w:eastAsia'), n)
    run.font.size = Pt(s); run.bold = b
    if c: run.font.color.rgb = RGBColor(*c)

def ah(doc, t, l=1):
    h = doc.add_heading(t, level=l)
    for r in h.runs:
        ss = {0:16,1:14,2:13,3:12}
        r.font.size = Pt(ss.get(l,12)); r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def ap(doc, t, indent=True):
    if not t.strip(): return
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
    if indent and not t.startswith('!') and not t.startswith('|') and not t.startswith('```'):
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(t); sf(r)

def ai(doc, ip, cap=''):
    if not os.path.exists(ip):
        alt = os.path.join(IMG, os.path.basename(ip))
        if os.path.exists(alt): ip = alt
        else: return
    try:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(); r.add_picture(ip, width=Inches(5.5))
        if cap:
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(10)
            r = cp.add_run(cap); sf(r, name='宋体', s=9)
    except: pass

with open(MD, 'r', encoding='utf-8') as f: content = f.read()
doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17); sec.right_margin = Cm(3.17)
style = doc.styles['Normal']; style.font.name = '宋体'; style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

lines = content.split('\n'); i = 0; ic = 0
while i < len(lines):
    line = lines[i]; s = line.strip(); i += 1
    if not s: continue
    if s.startswith('# ') and not s.startswith('## '): ah(doc, s[2:], 1); continue
    if s.startswith('## ') and not s.startswith('### '): ah(doc, s[3:], 2); continue
    if s.startswith('### ') and not s.startswith('#### '): ah(doc, s[4:], 2); continue
    if s.startswith('#### '): ah(doc, s[5:], 3); continue
    m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', s)
    if m:
        alt, rel = m.group(1), m.group(2)
        apath = rel if os.path.isabs(rel) else os.path.join(BASE, rel)
        ai(doc, apath.replace('/', '\\'), alt if alt else ''); ic += 1; continue
    if s.startswith('|'):
        tl = [s]
        while i < len(lines):
            nl = lines[i].strip()
            if nl.startswith('|'): tl.append(nl); i += 1
            else: break
        dl = [l for l in tl if not re.match(r'^\|[\s\-\|:]+\|$', l)]
        if dl:
            rows = [[c.strip() for c in l.split('|')[1:-1]] for l in dl]
            mc = max(len(r) for r in rows)
            t = doc.add_table(rows=len(rows), cols=mc, style='Table Grid')
            for ri, rd in enumerate(rows):
                for ci in range(mc):
                    ct = rd[ci] if ci < len(rd) else ''
                    t.cell(ri, ci).text = ct.strip()
                    for pp in t.cell(ri, ci).paragraphs:
                        for rr in pp.runs: sf(rr, s=9, b=(ri==0))
            doc.add_paragraph(); continue
    if s.startswith('```'):
        while i < len(lines) and not lines[i].strip().startswith('```'): i += 1
        if i < len(lines): i += 1; continue
    if re.match(r'^[-*+]\s', s):
        items = [re.sub(r'^[-*+]\s+', '', s)]
        while i < len(lines) and re.match(r'^[-*+]\s', lines[i].strip()):
            items.append(re.sub(r'^[-*+]\s+', '', lines[i].strip())); i += 1
        for item in items:
            p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5; p.paragraph_format.left_indent = Cm(0.5)
            r = p.add_run('• ' + item); sf(r)
        continue
    if re.match(r'^\d+[\.\、）)]\s', s):
        items = [re.sub(r'^\d+[\.\、）)]\s+', '', s)]
        while i < len(lines) and re.match(r'^\d+[\.\、）)]\s', lines[i].strip()):
            items.append(re.sub(r'^\d+[\.\、）)]\s+', '', lines[i].strip())); i += 1
        for idx, item in enumerate(items, 1):
            p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run(f'{idx}. {item}'); sf(r)
        continue
    if s == '---' or s == '***': doc.add_page_break(); continue
    if '**' in s:
        p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0.74)
        for part in re.split(r'(\*\*[^*]+\*\*)', s):
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2]); r.bold = True
            else: r = p.add_run(part)
            sf(r)
    else: ap(doc, s)

doc.save(DOC)
print(f"DOC: {DOC} ({ic} images)")
try:
    import win32com.client
    w = win32com.client.Dispatch("Word.Application"); w.Visible = False
    d = w.Documents.Open(DOC); d.SaveAs(PDF, FileFormat=17)
    d.Close(); w.Quit()
    print(f"PDF: {PDF}")
except Exception as e:
    print(f"PDF skip: {e}")
